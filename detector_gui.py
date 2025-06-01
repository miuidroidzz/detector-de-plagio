import tkinter as tk
from tkinter import messagebox, filedialog
from tkinter import ttk
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from fpdf import FPDF
from docx import Document

def verificar_plagio():
    texto1 = entrada_texto1.get("1.0", tk.END).strip()
    texto2 = entrada_texto2.get("1.0", tk.END).strip()
    
    if not texto1 or not texto2:
        messagebox.showwarning("Aviso", "Por favor, insira os dois textos.")
        return

    vetorizar = TfidfVectorizer().fit_transform([texto1, texto2])
    matriz_similaridade = cosine_similarity(vetorizar[0:1], vetorizar[1:2])
    similaridade = matriz_similaridade[0][0] * 100

    resultado_var.set(f"Similaridade: {similaridade:.2f}%")

def exportar_pdf():
    if not resultado_var.get():
        messagebox.showwarning("Aviso", "Realize a verificação antes de exportar.")
        return
    caminho = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
    if caminho:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, f"Resultado da Verificação:\n{resultado_var.get()}")
        pdf.output(caminho)
        messagebox.showinfo("Sucesso", f"Resultado exportado como PDF: {caminho}")

def exportar_docx():
    if not resultado_var.get():
        messagebox.showwarning("Aviso", "Realize a verificação antes de exportar.")
        return
    caminho = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if caminho:
        doc = Document()
        doc.add_heading("Resultado da Verificação", level=1)
        doc.add_paragraph(resultado_var.get())
        doc.save(caminho)
        messagebox.showinfo("Sucesso", f"Resultado exportado como DOCX: {caminho}")

janela = tk.Tk()
janela.title("Detector de Plágio")
janela.geometry("700x600")
janela.resizable(False, False)

style = ttk.Style()
style.theme_use('clam')

ttk.Label(janela, text="Texto 1:").pack(pady=5)
entrada_texto1 = tk.Text(janela, height=10, width=80, relief="solid", bd=1)
entrada_texto1.pack(pady=5)

ttk.Label(janela, text="Texto 2:").pack(pady=5)
entrada_texto2 = tk.Text(janela, height=10, width=80, relief="solid", bd=1)
entrada_texto2.pack(pady=5)

btn_verificar = ttk.Button(janela, text="Verificar Plágio", command=verificar_plagio)
btn_verificar.pack(pady=15)

resultado_var = tk.StringVar()
resultado_label = ttk.Label(janela, textvariable=resultado_var, font=("Arial", 14))
resultado_label.pack(pady=10)

frame_export = ttk.Frame(janela)
frame_export.pack(pady=10)

btn_exportar_pdf = ttk.Button(frame_export, text="Exportar PDF", command=exportar_pdf)
btn_exportar_pdf.grid(row=0, column=0, padx=10)

btn_exportar_docx = ttk.Button(frame_export, text="Exportar DOCX", command=exportar_docx)
btn_exportar_docx.grid(row=0, column=1, padx=10)

janela.mainloop()